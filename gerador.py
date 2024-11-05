"""
desenvolvido por: Leonardo Medeiros

lib usada: python-docx

"""

import random
from docx import Document

def inicio():
    print(30 * '-')
    print(' ')
    print(' Bem vindo ao Gerador de provas')
    print(' ')
    print(' DESENVOLVIDO POR: Leonardo Medeiros')
    print(' ')
    print(30 * '-')

# le as questoes do arquivo txt
def ler_questoes():
    with open('db_questoes.txt', 'r', encoding='utf-8') as questions_file:
        lines = questions_file.readlines()
      
        lines = [line.strip() for line in lines if line.strip()]

        questoes = [lines[i:i+5] for i in range(0, len(lines), 5)]
        return questoes

def gerador_de_prova(n_prova):
 
    questoes = ler_questoes()

    # escolhe 10 questões aletoarias
    perguntas_aleatorias = random.sample(questoes, 10)
    
    # Criar um documento Word
    nome_arquivo = f'prova_historia-{n_prova}.docx'
    document = Document()
    
    # Adicionar informações ao documento
    document.add_heading('Prova de História', level=1)
    document.add_paragraph('Nome: ____________________________________')
    document.add_paragraph('Data: ____________________________________')
    document.add_paragraph('Turma: ___________________________________')
    
   
    for i, pergunta in enumerate(perguntas_aleatorias, start=1):
        document.add_paragraph(f"{str(i)} - {pergunta[0]}")  
        for j in range(1, len(pergunta)):
            document.add_paragraph(pergunta[j], style='List Bullet')  
        document.add_paragraph('')  

   
    document.save(nome_arquivo)
    
    print(f"Perguntas salvas no arquivo {nome_arquivo}")

def main():
    inicio()
    n = int(input('Digite a quantidade de provas que você deseja gerar: '))
    for n_prova in range(1, n + 1):
        gerador_de_prova(n_prova)

if __name__ == '__main__':
    main()